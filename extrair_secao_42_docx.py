from docx import Document
import re
import json

def extrair_secao_42(docx_path):
    """
    Extrai especificamente a seção 4.2 do documento DOCX
    """
    try:
        doc = Document(docx_path)
        
        print("📄 Carregando documento DOCX...")
        print(f"Total de parágrafos: {len(doc.paragraphs)}")
        
        # Procurar pela seção 4.2
        secao_42_iniciada = False
        secao_42_finalizada = False
        paragrafos_secao_42 = []
        
        for i, para in enumerate(doc.paragraphs):
            texto = para.text.strip()
            
            # Detectar início da seção 4.2
            if re.match(r'4\.2\s+RESULTADOS', texto, re.IGNORECASE):
                secao_42_iniciada = True
                print(f"✅ Seção 4.2 encontrada no parágrafo {i}: {texto[:80]}...")
                paragrafos_secao_42.append(f"[PARÁGRAFO {i}] {texto}")
                continue
            
            # Detectar fim da seção 4.2 (início de uma nova seção principal)
            if secao_42_iniciada and re.match(r'5\.\d+|CAPÍTULO\s+5|5\s+DISCUSSÃO', texto, re.IGNORECASE):
                secao_42_finalizada = True
                print(f"🔚 Fim da seção 4.2 detectado no parágrafo {i}: {texto[:80]}...")
                break
            
            # Coletar parágrafos da seção 4.2
            if secao_42_iniciada and not secao_42_finalizada:
                if texto:  # Só adicionar parágrafos não vazios
                    paragrafos_secao_42.append(f"[PARÁGRAFO {i}] {texto}")
        
        print(f"📊 Parágrafos extraídos da seção 4.2: {len(paragrafos_secao_42)}")
        
        # Salvar seção extraída
        secao_42_texto = '\n\n'.join(paragrafos_secao_42)
        with open('secao_4_2_extraida.txt', 'w', encoding='utf-8') as f:
            f.write(secao_42_texto)
        
        print("💾 Seção 4.2 salva em 'secao_4_2_extraida.txt'")
        
        # Procurar especificamente por dados que precisam ser corrigidos
        encontrar_dados_para_correcao(paragrafos_secao_42)
        
        return paragrafos_secao_42
        
    except Exception as e:
        print(f"❌ Erro ao processar documento: {e}")
        return None

def encontrar_dados_para_correcao(paragrafos):
    """
    Localiza os dados específicos que precisam ser corrigidos
    """
    print("\n🔍 PROCURANDO DADOS PARA CORREÇÃO...")
    
    dados_encontrados = {
        'pilares_csc': [],
        'ods_valores': [],
        'tecnologias': [],
        'paragrafos_relevantes': []
    }
    
    for i, para in enumerate(paragrafos):
        texto = para.lower()
        
        # Procurar por pilares CSC
        if 'tecnologia e inovação' in texto and 'menções' in texto:
            dados_encontrados['pilares_csc'].append((i, para))
            print(f"📍 Pilar CSC encontrado: {para[:100]}...")
        
        if 'mobilidade' in texto and 'menções' in texto:
            dados_encontrados['pilares_csc'].append((i, para))
            print(f"📍 Pilar Mobilidade encontrado: {para[:100]}...")
        
        if 'governança' in texto and 'menções' in texto:
            dados_encontrados['pilares_csc'].append((i, para))
            print(f"📍 Pilar Governança encontrado: {para[:100]}...")
        
        # Procurar por ODS específicos
        if re.search(r'ods\s+(7|9|13)', texto):
            dados_encontrados['ods_valores'].append((i, para))
            print(f"📍 ODS encontrado: {para[:100]}...")
        
        # Procurar por tecnologias
        if any(tech in texto for tech in ['iot', 'inteligência artificial', 'big data', 'gêmeos digitais']):
            dados_encontrados['tecnologias'].append((i, para))
            print(f"📍 Tecnologia encontrada: {para[:100]}...")
    
    # Salvar dados encontrados
    with open('dados_para_correcao.json', 'w', encoding='utf-8') as f:
        json.dump({
            'total_paragrafos': len(paragrafos),
            'pilares_csc_encontrados': len(dados_encontrados['pilares_csc']),
            'ods_encontrados': len(dados_encontrados['ods_valores']),
            'tecnologias_encontradas': len(dados_encontrados['tecnologias'])
        }, f, indent=2, ensure_ascii=False)
    
    print(f"\n📊 RESUMO DOS DADOS ENCONTRADOS:")
    print(f"   - Pilares CSC: {len(dados_encontrados['pilares_csc'])} ocorrências")
    print(f"   - ODS: {len(dados_encontrados['ods_valores'])} ocorrências")
    print(f"   - Tecnologias: {len(dados_encontrados['tecnologias'])} ocorrências")
    
    return dados_encontrados

if __name__ == "__main__":
    import glob
    docx_files = glob.glob('*.docx')
    if docx_files:
        docx_path = docx_files[0]
        print(f"📄 Processando arquivo: {docx_path}")
        paragrafos = extrair_secao_42(docx_path)
    else:
        print("❌ Nenhum arquivo DOCX encontrado")
        paragrafos = None
    
    if paragrafos:
        print("\n✅ EXTRAÇÃO CONCLUÍDA COM SUCESSO")
        print("📁 Arquivos gerados:")
        print("   - secao_4_2_extraida.txt")
        print("   - dados_para_correcao.json")
    else:
        print("❌ FALHA NA EXTRAÇÃO")
