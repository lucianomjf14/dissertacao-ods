from docx import Document
import re

def extrair_secao_42_corrigida():
    """
    Extrai a seção 4.2 do documento corrigido para verificação
    """
    try:
        doc_corrigido = "DISSERTAÇÃO versão final - Revisão Luciano 13_06_CORRIGIDO.docx"
        doc = Document(doc_corrigido)
        
        print("📄 Verificando documento corrigido...")
        print(f"Total de parágrafos: {len(doc.paragraphs)}")
        
        # Procurar pela seção 4.2
        secao_42_iniciada = False
        secao_42_finalizada = False
        paragrafos_secao_42 = []
        
        for i, para in enumerate(doc.paragraphs):
            texto = para.text.strip()
            
            # Detectar início da seção 4.2
            if re.match(r'4\.2\s+RESULTADOS', texto, re.IGNORECASE):
                secao_42_iniciada = True
                print(f"✅ Seção 4.2 encontrada no parágrafo {i}")
                paragrafos_secao_42.append(f"[PARÁGRAFO {i}] {texto}")
                continue
            
            # Detectar fim da seção 4.2
            if secao_42_iniciada and re.match(r'5\.\d+|CAPÍTULO\s+5|5\s+DISCUSSÃO', texto, re.IGNORECASE):
                secao_42_finalizada = True
                print(f"🔚 Fim da seção 4.2 detectado no parágrafo {i}")
                break
            
            # Coletar parágrafos da seção 4.2
            if secao_42_iniciada and not secao_42_finalizada:
                if texto:
                    paragrafos_secao_42.append(f"[PARÁGRAFO {i}] {texto}")
        
        # Salvar seção corrigida
        secao_42_texto = '\n\n'.join(paragrafos_secao_42)
        with open('secao_4_2_CORRIGIDA.txt', 'w', encoding='utf-8') as f:
            f.write(secao_42_texto)
        
        print(f"📊 Parágrafos extraídos: {len(paragrafos_secao_42)}")
        print("💾 Seção corrigida salva em 'secao_4_2_CORRIGIDA.txt'")
        
        # Verificar especificamente os pilares CSC
        print("\\n🔍 VERIFICANDO PILARES CSC CORRIGIDOS...")
        for para in paragrafos_secao_42:
            if any(termo in para.lower() for termo in ['mobilidade (38', 'tecnologia e inovação (38', 'governança (37', 'segurança (35']):
                print(f"✅ Pilar corrigido encontrado: {para[:100]}...")
        
        return True
        
    except Exception as e:
        print(f"❌ Erro: {e}")
        return False

if __name__ == "__main__":
    extrair_secao_42_corrigida()
