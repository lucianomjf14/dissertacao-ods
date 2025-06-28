from docx import Document
import pandas as pd
import re
from collections import Counter
import shutil
import os

def fazer_backup_e_corrigir():
    """
    Faz backup do documento original e aplica correções diretamente
    """
    
    import glob
    docx_files = glob.glob('*.docx')
    if not docx_files:
        print("❌ Nenhum arquivo DOCX encontrado")
        return False
    
    docx_original = docx_files[0]
    print(f"📄 Documento original: {docx_original}")
    
    # Fazer backup
    backup_name = docx_original.replace('.docx', '_BACKUP.docx')
    shutil.copy2(docx_original, backup_name)
    print(f"💾 Backup criado: {backup_name}")
    
    # Carregar dados corretos da planilha
    df = pd.read_excel("Análise de Conteúdo - Luciano 22_06.xlsx")
    
    # Contar ODS corretos
    ods_counter = Counter()
    for idx, row in df.iterrows():
        ods_raw = str(row['ODS Mencionados pelos Autores'])
        if ods_raw.lower() != 'geral' and ods_raw != 'nan':
            ods_encontrados = re.findall(r'ODS\s+(\d+)', ods_raw)
            for ods_num in ods_encontrados:
                ods_counter[int(ods_num)] += 1
    
    total_registros = len(df)
    
    print(f"\\n📊 DADOS CORRETOS DA PLANILHA:")
    print(f"   - ODS 7: {ods_counter[7]} menções ({ods_counter[7]/total_registros*100:.2f}%)")
    print(f"   - ODS 9: {ods_counter[9]} menções ({ods_counter[9]/total_registros*100:.2f}%)")
    print(f"   - ODS 11: {ods_counter[11]} menções ({ods_counter[11]/total_registros*100:.2f}%)")
    print(f"   - ODS 13: {ods_counter[13]} menções ({ods_counter[13]/total_registros*100:.2f}%)")
    
    # Abrir documento original para edição
    try:
        doc = Document(docx_original)
        print(f"📄 Documento carregado com {len(doc.paragraphs)} parágrafos")
        
        correções_realizadas = 0
        
        # CORREÇÕES ESPECÍFICAS
        for i, para in enumerate(doc.paragraphs):
            texto_original = para.text.strip()
            texto_modificado = texto_original
            
            # Correção 1: Tecnologia e Inovação
            if "Tecnologia e Inovação (32 menções, presente em 85% dos estudos)" in texto_original:
                texto_modificado = texto_original.replace(
                    "Tecnologia e Inovação (32 menções, presente em 85% dos estudos)",
                    "Tecnologia e Inovação (38 menções, presente em 92,7% dos estudos)"
                )
                correções_realizadas += 1
                print(f"✅ {correções_realizadas}. Corrigido Tecnologia e Inovação (parágrafo {i})")
            
            # Correção 2: Governança  
            elif "Governança (30 menções, 73% dos estudos)" in texto_original:
                texto_modificado = texto_original.replace(
                    "Governança (30 menções, 73% dos estudos)",
                    "Governança (37 menções, 90,2% dos estudos)"
                )
                correções_realizadas += 1
                print(f"✅ {correções_realizadas}. Corrigido Governança (parágrafo {i})")
            
            # Correção 3: Meio Ambiente
            elif "Meio Ambiente (27 menções, 63% dos estudos)" in texto_original:
                texto_modificado = texto_original.replace(
                    "Meio Ambiente (27 menções, 63% dos estudos)",
                    "Meio Ambiente (33 menções, 80,5% dos estudos)"
                )
                correções_realizadas += 1
                print(f"✅ {correções_realizadas}. Corrigido Meio Ambiente (parágrafo {i})")
            
            # Correção 4: Energia
            elif "Energia (23 menções, 59% dos estudos)" in texto_original:
                texto_modificado = texto_original.replace(
                    "Energia (23 menções, 59% dos estudos)",
                    "Energia (34 menções, 82,9% dos estudos)"
                )
                correções_realizadas += 1
                print(f"✅ {correções_realizadas}. Corrigido Energia (parágrafo {i})")
            
            # Correção 5: Urbanismo
            elif "Urbanismo (19 menções, 46% dos estudos)" in texto_original:
                texto_modificado = texto_original.replace(
                    "Urbanismo (19 menções, 46% dos estudos)",
                    "Urbanismo (23 menções, 56,1% dos estudos)"
                )
                correções_realizadas += 1
                print(f"✅ {correções_realizadas}. Corrigido Urbanismo (parágrafo {i})")
            
            # Correção 6: Mobilidade (mais complexa)
            elif "Mobilidade (18 menções, 44% dos estudos)" in texto_original:
                texto_modificado = texto_original.replace(
                    "Mobilidade (18 menções, 44% dos estudos): o sexto pilar mais frequente",
                    "Mobilidade (38 menções, 92,7% dos estudos): empatado como o pilar mais frequente"
                )
                correções_realizadas += 1
                print(f"✅ {correções_realizadas}. Corrigido Mobilidade (parágrafo {i})")
            
            # Correção 7: Reordenar descrição dos pilares
            elif "este pilar apresenta a maior incidência, reforçando seu papel central" in texto_original and "Tecnologia e Inovação" in texto_original:
                texto_modificado = texto_original.replace(
                    "este pilar apresenta a maior incidência, reforçando seu papel central",
                    "empatado como o pilar mais frequente junto com Mobilidade, reforça seu papel central"
                )
                correções_realizadas += 1
                print(f"✅ {correções_realizadas}. Corrigida descrição TIC (parágrafo {i})")
            
            # Correção 8: Conclusão dos pilares
            elif "A predominância de Tecnologia e Inovação reflete" in texto_original:
                texto_modificado = texto_original.replace(
                    "A predominância de Tecnologia e Inovação reflete o papel fundamental das soluções digitais",
                    "O empate entre Mobilidade e Tecnologia e Inovação como pilares mais importantes reflete tanto a centralidade dos sistemas de transporte quanto o papel fundamental das soluções digitais"
                )
                correções_realizadas += 1
                print(f"✅ {correções_realizadas}. Corrigida conclusão pilares (parágrafo {i})")
            
            # Aplicar modificação se houve mudança
            if texto_modificado != texto_original:
                para.clear()
                para.add_run(texto_modificado)
        
        # Salvar documento modificado
        doc.save(docx_original)
        print(f"\\n💾 Documento original atualizado com {correções_realizadas} correções")
        
        return True
        
    except Exception as e:
        print(f"❌ Erro ao processar documento: {e}")
        return False

def adicionar_pilar_seguranca():
    """
    Adiciona o pilar Segurança que estava faltando
    """
    import glob
    docx_files = glob.glob('*.docx')
    if not docx_files:
        return False
    
    doc = Document(docx_files[0])
    
    # Procurar onde inserir o pilar Segurança
    for i, para in enumerate(doc.paragraphs):
        texto = para.text.strip()
        
        # Inserir após Governança e antes de Energia
        if "Governança (37 menções, 90,2% dos estudos)" in texto:
            print(f"📍 Encontrado local para inserir Segurança após parágrafo {i}")
            
            # Criar novo parágrafo para Segurança
            novo_para = doc.add_paragraph()
            texto_seguranca = "4. Segurança (35 menções, 85,4% dos estudos): este pilar emerge como uma dimensão fundamental das cidades inteligentes, evidenciando a importância dos sistemas de monitoramento, videovigilância inteligente e gestão de emergências nas iniciativas urbanas. A presença significativa da Segurança nos estudos demonstra que as C.I. devem contemplar não apenas eficiência e sustentabilidade, mas também proteção e resiliência urbana."
            novo_para.add_run(texto_seguranca)
            
            print("✅ Pilar Segurança adicionado")
            
            # Salvar
            doc.save(docx_files[0])
            return True
    
    return False

if __name__ == "__main__":
    print("🔧 INICIANDO CORREÇÕES DIRETAS NO DOCUMENTO FINAL")
    print("="*60)
    
    sucesso = fazer_backup_e_corrigir()
    
    if sucesso:
        print("\\n✅ CORREÇÕES APLICADAS COM SUCESSO!")
        print("\\n🎯 PRÓXIMO: Executar auditoria para verificar precisão")
        
        # Executar auditoria imediatamente
        print("\\n🔍 EXECUTANDO AUDITORIA DE VERIFICAÇÃO...")
        
        import subprocess
        try:
            resultado = subprocess.run([
                "c:/Users/lucia/Dissertação/.venv/Scripts/python.exe", 
                "auditoria_dados_dissertacao.py"
            ], capture_output=True, text=True, cwd=".")
            
            print("📊 RESULTADO DA AUDITORIA:")
            print(resultado.stdout[-1000:])  # Últimas 1000 caracteres
            
        except Exception as e:
            print(f"⚠️ Erro na auditoria: {e}")
    else:
        print("❌ FALHA NAS CORREÇÕES")
