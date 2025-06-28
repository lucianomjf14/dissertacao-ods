#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script para corrigir valores dos ODS no documento final da dissertação
"""

import os
import shutil
from datetime import datetime
from docx import Document

def aplicar_correcoes_ods(doc_path):
    """Aplica correções específicas dos valores dos ODS no documento"""
    
    # Backup do documento
    backup_path = doc_path.replace('.docx', f'_backup_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx')
    shutil.copy2(doc_path, backup_path)
    print(f"📋 Backup criado: {backup_path}")
    
    # Abrir documento
    doc = Document(doc_path)
    
    correções_aplicadas = []
    
    # Definir correções específicas dos ODS
    correções_ods = [
        # ODS 11 - Corrigir valores
        ("20 estudos (48,78%)", "22 estudos (53,66%)"),
        ("20 artigos (48,78%)", "22 artigos (53,66%)"),
        ("48,78%", "53,66%"),
        ("48.78%", "53.66%"),
        
        # ODS 7 - Corrigir valores  
        ("9 estudos (21,95%)", "10 estudos (24,39%)"),
        ("9 artigos (21,95%)", "10 artigos (24,39%)"),
        ("21,95%", "24,39%"),
        ("21.95%", "24.39%"),
        
        # ODS 9 - Corrigir valores
        ("8 estudos (19,51%)", "11 estudos (26,83%)"),
        ("8 artigos (19,51%)", "11 artigos (26,83%)"),
        ("19,51%", "26,83%"),
        ("19.51%", "26.83%"),
        
        # ODS 13 - Corrigir valores
        ("8 estudos (19,51%)", "9 estudos (21,95%)"),
        ("8 artigos (19,51%)", "9 artigos (21,95%)"),
        
        # Remover ODS 1 e 14 (substituir por texto genérico ou remover)
        ("ODS 1", ""),
        ("ODS 14", ""),
        ("1 estudo (2,44%)", ""),
        ("1 artigo (2,44%)", ""),
    ]
    
    # Aplicar correções em parágrafos
    for paragrafo in doc.paragraphs:
        texto_original = paragrafo.text
        texto_novo = texto_original
        
        for busca, substituicao in correções_ods:
            if busca in texto_novo:
                texto_novo = texto_novo.replace(busca, substituicao)
                if texto_novo != texto_original:
                    correções_aplicadas.append(f"Parágrafo: '{busca}' → '{substituicao}'")
        
        if texto_novo != texto_original:
            # Limpar e reescrever parágrafo
            for run in paragrafo.runs:
                run.clear()
            paragrafo.add_run(texto_novo)
    
    # Aplicar correções em tabelas
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for paragrafo in celula.paragraphs:
                    texto_original = paragrafo.text
                    texto_novo = texto_original
                    
                    for busca, substituicao in correções_ods:
                        if busca in texto_novo:
                            texto_novo = texto_novo.replace(busca, substituicao)
                            if texto_novo != texto_original:
                                correções_aplicadas.append(f"Tabela: '{busca}' → '{substituicao}'")
                    
                    if texto_novo != texto_original:
                        # Limpar e reescrever parágrafo da célula
                        for run in paragrafo.runs:
                            run.clear()
                        paragrafo.add_run(texto_novo)
    
    # Salvar documento
    doc.save(doc_path)
    
    return correções_aplicadas

def main():
    """Função principal"""
    doc_path = r"c:\Users\lucia\Dissertação\DISSERTAÇÃO versão final - Revisão Luciano 13_06.docx"
    
    if not os.path.exists(doc_path):
        print(f"❌ Documento não encontrado: {doc_path}")
        return
    
    print("🔧 INICIANDO CORREÇÃO DOS ODS")
    print("="*60)
    
    try:
        correções = aplicar_correcoes_ods(doc_path)
        
        print("✅ CORREÇÕES DOS ODS APLICADAS COM SUCESSO!")
        print(f"📊 Total de correções: {len(correções)}")
        
        if correções:
            print("\n🔍 DETALHES DAS CORREÇÕES:")
            for i, correcao in enumerate(correções[:20], 1):  # Mostrar primeiras 20
                print(f"   {i}. {correcao}")
            
            if len(correções) > 20:
                print(f"   ... e mais {len(correções) - 20} correções")
        
        print(f"\n💾 Documento atualizado: {doc_path}")
        
    except Exception as e:
        print(f"❌ Erro durante a correção: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()
